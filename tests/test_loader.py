"""生产级 Loader 测试：多格式 + 编码容错 + 错误隔离 + 递归。"""

import pytest

from health_rag.ingestion.loader import (
    load_document,
    load_documents_from_directory,
    load_text_file,
)


def _make_minimal_pdf(text: str, path) -> None:
    """构造一个最小但合法、含文本内容的 PDF 文件。"""
    content = f"BT /F1 24 Tf 100 700 Td ({text}) Tj ET".encode("ascii")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length "
        + str(len(content)).encode("ascii")
        + b" >>\nstream\n"
        + content
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    data = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(data))
        data.extend(f"{i} 0 obj\n".encode("ascii"))
        data.extend(obj)
        data.extend(b"\nendobj\n")

    xref_pos = len(data)
    data.extend(b"xref\n0 6\n")
    data.extend(b"0000000000 65535 f \n")
    for off in offsets:
        data.extend(f"{off:010d} 00000 n \n".encode("ascii"))
    data.extend(
        f"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode(
            "ascii"
        )
    )

    path.write_bytes(bytes(data))


def test_load_text_file(tmp_path):
    """测试加载单个文本文件（自包含，不依赖外部数据文件）。"""
    test_file = tmp_path / "test_health.txt"
    test_file.write_text(
        "健康饮食基础知识\n\n"
        "均衡饮食是维持人体健康的重要基础。\n",
        encoding="utf-8",
    )

    documents = load_text_file(test_file)

    assert len(documents) == 1
    assert "健康饮食基础知识" in documents[0].page_content
    assert documents[0].metadata["file_type"] == ".txt"


def test_load_markdown_file(tmp_path):
    """测试加载 Markdown 文件。"""
    test_file = tmp_path / "test_health.md"
    test_file.write_text(
        "# 健康饮食\n\n"
        "均衡饮食是维持健康的基础。\n",
        encoding="utf-8",
    )

    documents = load_text_file(test_file)

    assert len(documents) == 1
    assert "# 健康饮食" in documents[0].page_content
    assert documents[0].metadata["file_type"] == ".md"


def test_load_pdf(tmp_path):
    """测试加载 PDF 文件（逐页提取 + 页码元数据）。"""
    pdf_file = tmp_path / "test.pdf"
    _make_minimal_pdf("Healthy Diet Guide", pdf_file)

    documents = load_document(pdf_file)

    assert len(documents) == 1
    assert "Healthy Diet Guide" in documents[0].page_content
    assert documents[0].metadata["file_type"] == ".pdf"
    assert documents[0].metadata["page"] == 1


def test_load_docx(tmp_path):
    """测试加载 DOCX 文件。"""
    from docx import Document as DocxDocument

    docx_file = tmp_path / "test.docx"
    doc = DocxDocument()
    doc.add_paragraph("健康饮食指南")
    doc.add_paragraph("均衡饮食是维持健康的基础。")
    doc.save(str(docx_file))

    documents = load_document(docx_file)

    assert len(documents) == 1
    assert "健康饮食指南" in documents[0].page_content
    assert "均衡饮食" in documents[0].page_content
    assert documents[0].metadata["file_type"] == ".docx"


def test_load_gbk_encoded_file(tmp_path):
    """测试 GBK 编码文件（编码自动检测兜底）。"""
    test_file = tmp_path / "gbk_test.txt"
    test_file.write_bytes("健康饮食基础知识".encode("gbk"))

    documents = load_document(test_file)

    assert len(documents) == 1
    assert "健康饮食基础知识" in documents[0].page_content


def test_load_file_not_found(tmp_path):
    """测试加载不存在的文件。"""
    with pytest.raises(FileNotFoundError):
        load_text_file(tmp_path / "not_exist.txt")


def test_load_unsupported_type(tmp_path):
    """测试不支持的扩展名。"""
    bad_file = tmp_path / "test.xyz"
    bad_file.write_text("fake content", encoding="utf-8")

    with pytest.raises(ValueError):
        load_text_file(bad_file)


def test_load_directory_error_isolation(tmp_path):
    """测试错误隔离：坏文件被跳过，好文件正常加载。"""
    good = tmp_path / "good.md"
    good.write_text("# 健康\n\n正常内容。\n", encoding="utf-8")

    broken = tmp_path / "broken.pdf"
    broken.write_bytes(b"this is definitely not a valid pdf")

    documents = load_documents_from_directory(tmp_path)

    assert len(documents) == 1
    assert "正常内容" in documents[0].page_content


def test_load_directory_recursive(tmp_path):
    """测试递归加载子目录。"""
    (tmp_path / "sub").mkdir()

    root_file = tmp_path / "root.md"
    root_file.write_text("# 根目录\n", encoding="utf-8")

    sub_file = tmp_path / "sub" / "child.md"
    sub_file.write_text("# 子目录\n", encoding="utf-8")

    documents = load_documents_from_directory(tmp_path)

    assert len(documents) == 2
    assert any("根目录" in d.page_content for d in documents)
    assert any("子目录" in d.page_content for d in documents)