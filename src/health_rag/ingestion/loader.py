"""生产级文档加载器。

特性：
- 多格式支持：.txt / .md / .pdf / .docx
- 编码自动检测：utf-8-sig → utf-8 → gbk → latin-1 兜底（中文 GBK 文档不崩）
- 错误隔离：目录批量加载时，坏文件记录日志跳过，不中断整批
- 递归加载：支持子目录（按分类组织语料）
- 元数据：source / file_type / page（PDF 逐页）
- 输出 LangChain Document 对象（与下游 splitter / pipeline 兼容）
"""

import logging
from pathlib import Path

from langchain_core.documents import Document

logger = logging.getLogger(__name__)

SUPPORTED_SUFFIXES = {".txt", ".md", ".pdf", ".docx"}

# 编码检测级联：按顺序尝试，全部失败用 latin-1 兜底（永不失败）
_ENCODING_CASCADE = ("utf-8-sig", "utf-8", "gbk", "gb2312", "latin-1")


def _read_text_with_encoding(path: Path) -> str:
    """读取文本文件内容，自动检测编码。"""
    raw = path.read_bytes()

    for encoding in _ENCODING_CASCADE:
        try:
            return raw.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue

    # 理论不可达（latin-1 不抛错），防御性兜底
    return raw.decode("latin-1", errors="replace")


def _load_text(path: Path) -> list[Document]:
    """加载纯文本 / Markdown 文件。"""
    content = _read_text_with_encoding(path)

    return [
        Document(
            page_content=content,
            metadata={
                "source": str(path),
                "file_type": path.suffix.lower(),
            },
        )
    ]


def _load_pdf(path: Path) -> list[Document]:
    """加载 PDF 文件，逐页提取文本。"""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    documents = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": str(path),
                        "file_type": ".pdf",
                        "page": i + 1,
                    },
                )
            )

    return documents


def _load_docx(path: Path) -> list[Document]:
    """加载 DOCX 文件，按段落提取文本。"""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    return [
        Document(
            page_content="\n".join(parts),
            metadata={
                "source": str(path),
                "file_type": ".docx",
            },
        )
    ]


# 扩展名 → 加载函数 分发表
_LOADERS = {
    ".txt": _load_text,
    ".md": _load_text,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
}


def load_document(file_path: str | Path) -> list[Document]:
    """加载单个文档，按扩展名分发到对应加载器。

    Returns:
        文档列表（PDF 多页可能返回多个 Document）。

    Raises:
        FileNotFoundError: 文件不存在
        ValueError: 不支持的文件类型
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    loader = _LOADERS.get(path.suffix.lower())
    if loader is None:
        raise ValueError(
            f"Unsupported file type: {path.suffix}. "
            f"Supported: {', '.join(sorted(SUPPORTED_SUFFIXES))}"
        )

    return loader(path)


def load_text_file(file_path: str | Path) -> list[Document]:
    """Backward-compatible alias for load_document."""
    return load_document(file_path)


def load_documents_from_directory(
    directory: str | Path,
    recursive: bool = True,
) -> list[Document]:
    """递归加载目录下所有支持格式的文档。

    - 自动跳过不支持的扩展名
    - 坏文件：记录日志并跳过，不中断整批

    Returns:
        所有成功加载的 Document 对象。
    """
    dir_path = Path(directory)
    if not dir_path.is_dir():
        raise NotADirectoryError(f"Not a directory: {dir_path}")

    all_docs = []
    failed = []

    # 递归扫描（rglob）或仅当前目录（glob）
    iterator = dir_path.rglob("*") if recursive else dir_path.glob("*")

    for file_path in sorted(iterator):
        if not file_path.is_file():
            continue
        if file_path.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue

        try:
            all_docs.extend(load_document(file_path))
        except Exception as e:  # 错误隔离：单个文件失败不影响整批
            failed.append((str(file_path), str(e)))
            logger.warning("跳过文件 %s: %s", file_path, e)

    if failed:
        logger.warning(
            "共 %d 个文件加载失败: %s",
            len(failed),
            [f[0] for f in failed],
        )

    return all_docs